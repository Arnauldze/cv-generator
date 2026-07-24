from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document=Document()

# Profile picture
document.add_picture("girl.jpg", width=Inches(2.0))

# Name, phone and email details
name=input("Enter your name: ")
speak("Hello" + name + "How are you doing today?")

speak("Enter your phone number")
phone_number=input("Enter your phone number: ")
email=input("Enter your email: ")

document.add_paragraph(
    name + " | " + phone_number + " | " + email
)

# about me
document.add_heading("About me")
about_me=input("Tell about yourself ")
document.add_paragraph(about_me)

# work experience
document.add_heading("Work Experience")
p=document.add_paragraph()

company=input("Enter Company: ")
start_date=input("From Date: ")
finish_date=input("To Date: ")

p.add_run(company+" ").bold = True
p.add_run(start_date + "-" + finish_date +"\n").italic=True

experience_details = input("Describe your experience at " + company + " ")
p.add_run(experience_details)   

# more experiences
while True:
    has_more_experiences = input(
        "Do you have more experiences? Yes or No: "
    )
    if has_more_experiences.lower() == "yes":
        p=document.add_paragraph()

        company=input("Enter Company: ")
        start_date=input("From Date: ")
        finish_date=input("To Date: ")

        p.add_run(company+" ").bold = True
        p.add_run(start_date + "-" + finish_date +"\n").italic=True

        experience_details = input("Describe your experience at " + company + " ")
        p.add_run(experience_details)
    else:
        break    
        
# Skills
document.add_heading("Skills")
skill=input("Enter a skill: ")
p=document.add_paragraph(skill)

p.style="List Bullet"

while True:
     has_more_skills = input(
            "Do you have more skills? Yes or No: "
        )
     if has_more_skills.lower() == "yes":
         skill=input("Enter a skill: ")
         p=document.add_paragraph(skill)
         
         p.style="List Bullet"
     else:
         break
        
# footer
section = document.sections[0]
footer = section.footer
p=footer.paragraphs[0]
p.text="CV generated using pitchdev"

document.save("cv.docx") 